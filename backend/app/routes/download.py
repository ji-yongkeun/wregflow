from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.db.database import get_db
from app.models.analysis import AnalysisGroup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
import json
import os
import tempfile
from datetime import datetime
from io import BytesIO
from app.services.diagram_generator import (
    generate_process_mermaid,
    generate_interaction_mermaid,
    generate_decision_mermaid,
    mermaid_to_png,
    create_placeholder_png
)

router = APIRouter(prefix="/api/download", tags=["download"])

@router.get("/integration/{integration_id}/word")
async def download_integration_as_word(
    integration_id: str,
    db: Session = Depends(get_db)
):
    """
    통합 분석 결과를 Word 문서로 다운로드 (다이어그램 포함)
    """
    try:
        integration = db.query(AnalysisGroup).filter(
            AnalysisGroup.id == integration_id
        ).first()
        
        if not integration:
            raise HTTPException(status_code=404, detail="통합 분석을 찾을 수 없습니다")
        
        doc = Document()
        
        # 제목
        title = doc.add_heading(f'🎯 {integration.group_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 메타정보
        meta_para = doc.add_paragraph()
        meta_para.add_run('생성일시: ').bold = True
        meta_para.add_run(integration.created_at.strftime('%Y-%m-%d %H:%M:%S'))
        
        doc.add_paragraph()
        
        # ========== 다이어그램 섹션 ==========
        doc.add_heading('📊 통합 분석 다이어그램', level=1)
        
        data = integration.integrated_data or {}
        
        # 1. 통합 프로세스
        try:
            mermaid_code = generate_process_mermaid(data)
            if mermaid_code:
                doc.add_heading('📊 통합 프로세스 흐름', level=2)
                png_bytes = mermaid_to_png(mermaid_code)
                
                if png_bytes:
                    doc.add_picture(BytesIO(png_bytes), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph('(다이어그램 이미지 생성 실패)')
                
                doc.add_paragraph()
        except Exception as e:
            print(f"프로세스 다이어그램 생성 실패: {e}")
        
        # 2. 부서 간 데이터 흐름
        try:
            mermaid_code = generate_interaction_mermaid(data)
            if mermaid_code:
                doc.add_heading('🔗 부서 간 데이터 흐름', level=2)
                png_bytes = mermaid_to_png(mermaid_code)
                
                if png_bytes:
                    doc.add_picture(BytesIO(png_bytes), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph('(다이어그램 이미지 생성 실패)')
                
                doc.add_paragraph()
        except Exception as e:
            print(f"상호작용 다이어그램 생성 실패: {e}")
        
        # 3. 의사결정 포인트
        try:
            mermaid_code = generate_decision_mermaid(data)
            if mermaid_code:
                doc.add_heading('⚡ 의사결정 포인트', level=2)
                png_bytes = mermaid_to_png(mermaid_code)
                
                if png_bytes:
                    doc.add_picture(BytesIO(png_bytes), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph('(다이어그램 이미지 생성 실패)')
                
                doc.add_paragraph()
        except Exception as e:
            print(f"의사결정 다이어그램 생성 실패: {e}")
        
        # ========== 텍스트 섹션 ==========
        doc.add_heading('📋 상세 정보', level=1)
        
        if data:
            # 1. 통합 프로세스
            if 'integrated_process' in data:
                doc.add_heading('📊 통합 프로세스 흐름', level=2)
                process = data['integrated_process']
                if 'steps' in process:
                    for step in process['steps']:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"Step {step.get('step_number', '')} - {step.get('step_name', '')}").bold = True
                        p.add_run(f" ({step.get('responsible_department', '')})")
                doc.add_paragraph()
            
            # 2. 부서 간 데이터 흐름
            if 'department_interactions' in data:
                doc.add_heading('🔗 부서 간 데이터 흐름', level=2)
                interactions = data['department_interactions']
                for inter in interactions:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{inter.get('from_dept', '')} → {inter.get('to_dept', '')}").bold = True
                    p.add_run(f": {inter.get('data_flow', '')}")
                doc.add_paragraph()
            
            # 3. 의사결정 포인트
            if 'critical_decision_points' in data:
                doc.add_heading('⚡ 의사결정 포인트', level=2)
                decisions = data['critical_decision_points']
                for decision in decisions:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{decision.get('point_name', '')}").bold = True
                    p.add_run(f" ({decision.get('severity', 'Medium')} Impact)")
                    doc.add_paragraph(decision.get('description', ''), style='List Bullet 2')
                doc.add_paragraph()
            
            # 4. 리스크 포인트
            if 'risk_points' in data:
                doc.add_heading('⚠️ 리스크 포인트', level=2)
                risks = data['risk_points']
                for risk in risks:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{risk.get('point_name', '')} - {risk.get('severity', '')}").bold = True
                    doc.add_paragraph(f"완화방법: {risk.get('mitigation', '')}", style='List Bullet 2')
                doc.add_paragraph()
            
            # 5. 개선 기회
            if 'improvement_opportunities' in data:
                doc.add_heading('💡 개선 기회', level=2)
                improvements = data['improvement_opportunities']
                for improvement in improvements:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{improvement.get('area', '')} - {improvement.get('priority', 'Medium')}").bold = True
                    doc.add_paragraph(f"예상효과: {improvement.get('impact', '')}", style='List Bullet 2')
        
        filename = f"통합분석_{integration_id[:8]}.docx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        doc.save(filepath)
        
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word 다운로드 실패: {str(e)}")

@router.get("/integration/{integration_id}/excel")
async def download_integration_as_excel(
    integration_id: str,
    db: Session = Depends(get_db)
):
    """
    통합 분석 결과를 Excel로 다운로드 (다이어그램 포함)
    """
    try:
        from openpyxl.drawing.image import Image as XLImage
        
        integration = db.query(AnalysisGroup).filter(
            AnalysisGroup.id == integration_id
        ).first()
        
        if not integration:
            raise HTTPException(status_code=404, detail="통합 분석을 찾을 수 없습니다")
        
        wb = Workbook()
        
        # Sheet 1: 다이어그램
        ws_diagram = wb.active
        ws_diagram.title = "다이어그램"
        
        ws_diagram['A1'] = f"🎯 {integration.group_name} - 다이어그램"
        ws_diagram['A1'].font = ws_diagram['A1'].font.copy()
        
        data = integration.integrated_data or {}
        row = 3
        
        # 1. 통합 프로세스 다이어그램
        try:
            mermaid_code = generate_process_mermaid(data)
            if mermaid_code:
                ws_diagram[f'A{row}'] = '📊 통합 프로세스 흐름'
                row += 1
                
                png_bytes = mermaid_to_png(mermaid_code)
                if png_bytes:
                    img = XLImage(BytesIO(png_bytes))
                    img.width = 400
                    img.height = 200
                    ws_diagram.add_image(img, f'A{row}')
                    row += 15
        except Exception as e:
            print(f"프로세스 다이어그램 생성 실패: {e}")
        
        # 2. 부서 간 데이터 흐름 다이어그램
        try:
            mermaid_code = generate_interaction_mermaid(data)
            if mermaid_code:
                ws_diagram[f'A{row}'] = '🔗 부서 간 데이터 흐름'
                row += 1
                
                png_bytes = mermaid_to_png(mermaid_code)
                if png_bytes:
                    img = XLImage(BytesIO(png_bytes))
                    img.width = 400
                    img.height = 200
                    ws_diagram.add_image(img, f'A{row}')
                    row += 15
        except Exception as e:
            print(f"상호작용 다이어그램 생성 실패: {e}")
        
        # 3. 의사결정 포인트 다이어그램
        try:
            mermaid_code = generate_decision_mermaid(data)
            if mermaid_code:
                ws_diagram[f'A{row}'] = '⚡ 의사결정 포인트'
                row += 1
                
                png_bytes = mermaid_to_png(mermaid_code)
                if png_bytes:
                    img = XLImage(BytesIO(png_bytes))
                    img.width = 400
                    img.height = 200
                    ws_diagram.add_image(img, f'A{row}')
                    row += 15
        except Exception as e:
            print(f"의사결정 다이어그램 생성 실패: {e}")
        
        # Sheet 2: 의사결정 포인트 테이블
        ws_decisions = wb.create_sheet("의사결정")
        ws_decisions['A1'] = "⚡ 의사결정 포인트"
        
        row = 3
        ws_decisions[f'A{row}'] = "포인트"
        ws_decisions[f'B{row}'] = "심각도"
        ws_decisions[f'C{row}'] = "설명"
        row += 1
        
        if 'critical_decision_points' in data:
            for decision in data['critical_decision_points']:
                ws_decisions[f'A{row}'] = decision.get('point_name', '')
                ws_decisions[f'B{row}'] = decision.get('severity', '')
                ws_decisions[f'C{row}'] = decision.get('description', '')
                row += 1
        
        # Sheet 3: 리스크 포인트 테이블
        ws_risks = wb.create_sheet("리스크")
        ws_risks['A1'] = "⚠️ 리스크 포인트"
        
        row = 3
        ws_risks[f'A{row}'] = "리스크"
        ws_risks[f'B{row}'] = "심각도"
        ws_risks[f'C{row}'] = "완화방법"
        row += 1
        
        if 'risk_points' in data:
            for risk in data['risk_points']:
                ws_risks[f'A{row}'] = risk.get('point_name', '')
                ws_risks[f'B{row}'] = risk.get('severity', '')
                ws_risks[f'C{row}'] = risk.get('mitigation', '')
                row += 1
        
        # Sheet 4: 개선 기회 테이블
        ws_improvements = wb.create_sheet("개선")
        ws_improvements['A1'] = "💡 개선 기회"
        
        row = 3
        ws_improvements[f'A{row}'] = "영역"
        ws_improvements[f'B{row}'] = "우선순위"
        ws_improvements[f'C{row}'] = "예상효과"
        row += 1
        
        if 'improvement_opportunities' in data:
            for improvement in data['improvement_opportunities']:
                ws_improvements[f'A{row}'] = improvement.get('area', '')
                ws_improvements[f'B{row}'] = improvement.get('priority', '')
                ws_improvements[f'C{row}'] = improvement.get('impact', '')
                row += 1
        
        # 열 너비 조정
        for ws in [ws_diagram, ws_decisions, ws_risks, ws_improvements]:
            ws.column_dimensions['A'].width = 30
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 40
        
        filename = f"통합분석_{integration_id[:8]}.xlsx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        wb.save(filepath)
        
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=filename
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Excel 다운로드 실패: {str(e)}")

from pydantic import BaseModel
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

class ExcelExportPayload(BaseModel):
    process_name: str
    swim_lanes: list = []
    raci: list = []
    decisions: list = []
    system_interfaces: list = []

@router.post("/process-excel")
async def download_process_excel(
    payload: ExcelExportPayload,
    type: str = "fp"  # "fp" 단일 기준
):
    """
    3개의 JSON(Swimlane, RACI, Decisions) 및 인터페이스를 참고하여
    상세 프로세스 엑셀 파일(FP 기준)을 동적으로 생성 및 다운로드
    """
    try:
        wb = Workbook()
        ws = wb.active
        
        # 1. 시트 기본 정보 설정
        process_title = payload.process_name or "상세 업무 프로세스"
        
        # 2. 스타일 정의 (맑은 고딕 기준)
        title_font = Font(name="맑은 고딕", size=14, bold=True, color="FFFFFF")
        header_font = Font(name="맑은 고딕", size=10, bold=True, color="FFFFFF")
        data_font = Font(name="맑은 고딕", size=9)
        
        # FS는 Indigo/Blue 테마, FT는 Amber/Orange 테마
        if type == "ft":
            ws.title = "FT 시나리오 산정"
            theme_color = "B45309" # Amber
            title_text = f"⚡ {process_title} - FT(기능 테스트) 시나리오 산정"
        else:
            ws.title = "FS 기능정의 산정"
            theme_color = "3730A3" # Indigo
            title_text = f"📋 {process_title} - FS(기능 요건정의) 산정"
            
        title_fill = PatternFill(start_color=theme_color, end_color=theme_color, fill_type="solid")
        header_fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid") # Dark Slate
        
        # 테두리 및 정렬 설정
        thin_border = Border(
            left=Side(style='thin', color='CBD5E1'),
            right=Side(style='thin', color='CBD5E1'),
            top=Side(style='thin', color='CBD5E1'),
            bottom=Side(style='thin', color='CBD5E1')
        )
        
        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
        
        # 3. 타이틀 렌더링
        ws.merge_cells('A1:I1')
        ws['A1'] = title_text
        ws['A1'].font = title_font
        ws['A1'].fill = title_fill
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 40
        
        # 4. 데이터 플래트닝 및 정렬
        steps = []
        for lane in payload.swim_lanes:
            role = lane.get("role", "")
            for step_obj in lane.get("steps", []):
                steps.append({
                    "order": step_obj.get("order", 1),
                    "action": step_obj.get("action", step_obj.get("step_name", "")),
                    "decision": step_obj.get("decision", False),
                    "role": role,
                    "description": step_obj.get("description", ""),
                    "outputs": step_obj.get("outputs", []),
                    "related_editions": step_obj.get("related_editions", [])
                })
        steps.sort(key=lambda x: x["order"])
        
        # 5. 헤더 및 데이터 바인딩
        headers = []
        if type == "ft":
            headers = [
                "테스트 단계", "테스트 주체 (역할)", "테스트 시나리오명", "테스트 절차 (상세)", 
                "사전 조건 / 입력 데이터", "예상 결과 (Outputs)", "확인 시스템 (인터페이스)", 
                "의사결정 분기 결과", "관련 규정 편"
            ]
        else:
            headers = [
                "단계 번호", "담당 부서 / 주체", "업무 프로세스 단계명", "업무 상세 설명", 
                "RACI 역할 분담", "생성 산출물 (Outputs)", "연계 시스템", 
                "의사결정 분기 내용", "관련 규정 편"
            ]
            
        # 헤더 출력
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
            
        ws.row_dimensions[3].height = 28
        
        # 데이터 출력 및 매칭
        row_idx = 4
        for step in steps:
            action_clean = step["action"].strip()
            
            # RACI 매칭
            matched_raci = None
            for r in payload.raci:
                task_clean = r.get("task", "").strip()
                if task_clean and (task_clean == action_clean or task_clean in action_clean or action_clean in task_clean):
                    matched_raci = r
                    break
                    
            # Decisions 매칭
            matched_decision = None
            for d in payload.decisions:
                q_clean = d.get("question", d.get("point_name", "")).strip()
                if q_clean and (q_clean == action_clean or q_clean in action_clean or q_clean in step["description"]):
                    matched_decision = d
                    break
                    
            # Interfaces 매칭
            matched_interfaces = []
            for interface in payload.system_interfaces:
                sys_a = interface.get("system_a", "").lower()
                sys_b = interface.get("system_b", "").lower()
                desc_lower = (step["action"] + " " + step["description"]).lower()
                if (sys_a and sys_a in desc_lower) or (sys_b and sys_b in desc_lower):
                    matched_interfaces.append(f"{interface.get('system_a', '')} ↔ {interface.get('system_b', '')} ({interface.get('interface_type', '')})")
            
            # 출력 데이터 구성
            raci_text = ""
            if matched_raci:
                raci_text = f"R(수행): {matched_raci.get('responsible', '-')}\nA(책임): {matched_raci.get('accountable', '-')}\nC(합의): {matched_raci.get('consulted', '-')}\nI(보고): {matched_raci.get('informed', '-')}"
            
            decision_text = ""
            if matched_decision:
                decision_text = f"Yes: {matched_decision.get('yes_outcome', matched_decision.get('description', ''))}\nNo: {matched_decision.get('no_outcome', matched_decision.get('impact', ''))}"
            elif step["decision"]:
                decision_text = "의사결정 분기 분석 대상"
                
            outputs_text = "\n".join(step["outputs"]) if step["outputs"] else ""
            interfaces_text = "\n".join(matched_interfaces) if matched_interfaces else ""
            editions_text = ", ".join([f"{ed}편" for ed in step["related_editions"]]) if step["related_editions"] else ""
            
            # 셀 쓰기
            if type == "ft":
                # FT 기준
                ws.cell(row=row_idx, column=1, value=f"TS-{step['order']:02d}").alignment = center_align
                ws.cell(row=row_idx, column=2, value=step["role"]).alignment = center_align
                ws.cell(row=row_idx, column=3, value=step["action"]).alignment = left_align
                ws.cell(row=row_idx, column=4, value=step["description"] or "-").alignment = left_align
                ws.cell(row=row_idx, column=5, value="사전 심사 요건 만족 여부" if step["decision"] else "-").alignment = left_align
                ws.cell(row=row_idx, column=6, value=outputs_text or "처리 결과 반영").alignment = left_align
                ws.cell(row=row_idx, column=7, value=interfaces_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=8, value=decision_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=9, value=editions_text or "-").alignment = center_align
            else:
                # FS 기준
                ws.cell(row=row_idx, column=1, value=step["order"]).alignment = center_align
                ws.cell(row=row_idx, column=2, value=step["role"]).alignment = center_align
                ws.cell(row=row_idx, column=3, value=step["action"]).alignment = left_align
                ws.cell(row=row_idx, column=4, value=step["description"] or "-").alignment = left_align
                ws.cell(row=row_idx, column=5, value=raci_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=6, value=outputs_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=7, value=interfaces_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=8, value=decision_text or "-").alignment = left_align
                ws.cell(row=row_idx, column=9, value=editions_text or "-").alignment = center_align
                
            # 전체 셀 폰트 및 테두리 설정
            for col_idx in range(1, 10):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.font = data_font
                cell.border = thin_border
                
            ws.row_dimensions[row_idx].height = 42 # 충분히 높게 주어 텍스트 랩핑 수용
            row_idx += 1
            
        # 열 너비 설정
        col_widths = {
            1: 12, # 단계
            2: 20, # 부서
            3: 28, # 단계명
            4: 40, # 상세설명
            5: 25, # RACI or 사전조건
            6: 25, # 산출물 or 예상결과
            7: 25, # 연계시스템
            8: 30, # 의사결정 분기
            9: 12  # 관련 편
        }
        for col_idx, width in col_widths.items():
            ws.column_dimensions[get_column_letter(col_idx)].width = width
            
        # 임시 파일 저장 및 전송
        suffix_type = "FS" if type == "fs" else "FT"
        clean_title = "".join([c for c in process_title if c.isalnum() or c in (" ", "_", "-")]).strip()
        import urllib.parse
        encoded_filename = urllib.parse.quote(f"{clean_title}_{suffix_type}_{datetime.now().strftime('%m%d')}.xlsx")
        filepath = os.path.join(tempfile.gettempdir(), f"download_{suffix_type}_{row_idx}.xlsx")
        
        wb.save(filepath)
        
        # 파일 전송 및 삭제 대응
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=f"{clean_title}_{suffix_type}.xlsx"
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Excel 내보내기 실패: {str(e)}")
